import ollama
import os
import json
from rich.console import Console
from rich.markdown import Markdown
from duckduckgo_search import DDGS
import requests
from bs4 import BeautifulSoup
from readability import Document
import docx
import PyPDF2

console = Console()
HISTORY_FILE = "history.json"
DOCS_FOLDER = "docs"

# ------------------------
# Memory functions
# ------------------------
def load_history():
    if os.path.exists(HISTORY_FILE):
        with open(HISTORY_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return []

def save_history(history):
    with open(HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump(history, f, indent=2)

# ------------------------
# Web search function
# ------------------------
def web_search(query, max_results=3):
    """Search DuckDuckGo and return top results as readable summaries"""
    results = []
    with DDGS() as ddgs:
        for r in ddgs.text(query, max_results=max_results):
            url = r.get("href")
            if not url:
                continue
            try:
                resp = requests.get(url, timeout=10)
                doc = Document(resp.text)
                title = doc.title()
                summary = doc.summary()
                text = BeautifulSoup(summary, "lxml").get_text()
                results.append(f"{title}\n{text[:500]}... [source: {url}]")
            except Exception as e:
                results.append(f"Could not fetch {url} ({e})")
    return results

# ------------------------
# Document reading functions
# ------------------------
def read_text_file(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def read_docx_file(path):
    doc = docx.Document(path)
    return "\n".join([para.text for para in doc.paragraphs])

def read_pdf_file(path):
    text = ""
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() or ""
    return text

def load_documents(folder=DOCS_FOLDER):
    """Load all supported docs into a single string"""
    if not os.path.exists(folder):
        os.makedirs(folder)
        return "No documents found. Drop .txt, .pdf, or .docx files into the docs folder."

    texts = []
    for fn in os.listdir(folder):
        path = os.path.join(folder, fn)
        if fn.endswith(".txt"):
            texts.append(read_text_file(path))
        elif fn.endswith(".docx"):
            texts.append(read_docx_file(path))
        elif fn.endswith(".pdf"):
            texts.append(read_pdf_file(path))
    return "\n\n".join(texts) if texts else "No documents found in the docs folder."

# ------------------------
# Agent reply function
# ------------------------
def agent_reply(history, user_msg):
    user_msg_lower = user_msg.lower().strip()

    # Web search trigger
    if user_msg_lower.startswith("search:"):
        query = user_msg.split("search:",1)[1].strip()
        results = web_search(query)
        return "\n\n".join(results) if results else "No results found."

    # Document reading trigger
    if user_msg_lower.startswith("docs"):
        docs_text = load_documents()
        return f"Here’s what I found in your documents:\n\n{docs_text[:1000]}..."  # limit preview to 1000 chars

    # Normal chat with Ollama
    history.append({"role": "user", "content": user_msg})
    response = ollama.chat(
        model="llama3.1:8b",
        messages=history
    )
    reply = response["message"]["content"]
    history.append({"role": "assistant", "content": reply})
    return reply

# ------------------------
# Main loop
# ------------------------
def main():
    console.print("[bold green]AI Agent: Memory + Web Search + Document Reading![/bold green]")
    console.print("Type 'exit' to quit. For web search, type: search: <your query>")
    console.print("To read your documents, type: docs\n")

    history = load_history()

    while True:
        user_msg = input("\nYou: ").strip()
        if not user_msg:
            continue

        if user_msg.lower() in ["exit", "quit", "bye"]:
            console.print("[red]Goodbye![/red]")
            save_history(history)
            break

        with console.status("Thinking..."):
            out = agent_reply(history, user_msg)

        console.print(Markdown(out))
        save_history(history)

if __name__ == "__main__":
    main()
