# ai_agent_pro.py — Pro ChatGPT-like Agent (Streamlit + Gradio)
import os, io, re, json, time, math, pickle, traceback
from typing import List, Dict, Tuple

import numpy as np
import faiss

# Web search + article cleaning
import requests
from bs4 import BeautifulSoup
from readability import Document
try:
    from ddgs import DDGS         # preferred package name
except Exception:
    from duckduckgo_search import DDGS  # fallback name

# Files
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

# Embeddings (CPU to avoid meta-tensor error)
from sentence_transformers import SentenceTransformer

# LLM backends
DEFAULT_OLLAMA_MODEL = "llama3.1:8b"
OLLAMA_AVAILABLE = False
HF_AVAILABLE = False
try:
    import ollama
    OLLAMA_AVAILABLE = True
except Exception:
    OLLAMA_AVAILABLE = False

try:
    from transformers import pipeline
    hf_pipe = pipeline("text2text-generation", model="google/flan-t5-large")
    HF_AVAILABLE = True
except Exception:
    HF_AVAILABLE = False

# ---------- Paths / Config ----------
APP_TITLE = "Pro AI Agent — Chat + RAG + Web"
BASE_DIR = os.path.abspath(".")
DOCS_DIR = os.path.join(BASE_DIR, "docs")
os.makedirs(DOCS_DIR, exist_ok=True)

INDEX_FILE  = os.path.join(BASE_DIR, "faiss.index")
CHUNKS_FILE = os.path.join(BASE_DIR, "chunks.pkl")
METAS_FILE  = os.path.join(BASE_DIR, "metas.pkl")

EMBED_MODEL_NAME = "all-MiniLM-L6-v2"

def get_embedder():
    # Force CPU for stability on Windows/PyTorch (prevents meta-tensor error)
    return SentenceTransformer(EMBED_MODEL_NAME, device="cpu")

EMBEDDER = get_embedder()

# ---------- Chunking / Reading ----------
def chunk_text(text: str, chunk_size=1200, overlap=200) -> List[str]:
    text = text or ""
    chunks, n, start = [], len(text), 0
    while start < n:
        end = min(n, start + chunk_size)
        chunks.append(text[start:end])
        if end == n: break
        start = max(0, end - overlap)
    return chunks

def read_txt(b: bytes) -> str:
    return b.decode("utf-8", errors="ignore")

def read_pdf(b: bytes) -> str:
    buf = io.BytesIO(b)
    reader = PdfReader(buf)
    return "\n".join((p.extract_text() or "") for p in reader.pages)

def read_docx(b: bytes) -> str:
    buf = io.BytesIO(b)
    doc = DocxDocument(buf)
    return "\n".join(p.text for p in doc.paragraphs)

# ---------- RAG persistence ----------
def save_rag(index, chunks, metas):
    faiss.write_index(index, INDEX_FILE)
    with open(CHUNKS_FILE, "wb") as f: pickle.dump(chunks, f)
    with open(METAS_FILE, "wb") as f: pickle.dump(metas, f)

def load_rag():
    if not (os.path.exists(INDEX_FILE) and os.path.exists(CHUNKS_FILE) and os.path.exists(METAS_FILE)):
        return None, [], []
    try:
        index = faiss.read_index(INDEX_FILE)
        with open(CHUNKS_FILE, "rb") as f: chunks = pickle.load(f)
        with open(METAS_FILE, "rb") as f: metas  = pickle.load(f)
        return index, chunks, metas
    except Exception:
        return None, [], []

def rebuild_rag_from_docs() -> Tuple[faiss.Index, List[str], List[dict]]:
    all_chunks, metas = [], []
    for fn in os.listdir(DOCS_DIR):
        path = os.path.join(DOCS_DIR, fn)
        if not os.path.isfile(path): continue
        try:
            with open(path, "rb") as f: data = f.read()
            lo = fn.lower()
            if lo.endswith(".pdf"):  text = read_pdf(data)
            elif lo.endswith(".docx"): text = read_docx(data)
            elif lo.endswith(".txt"):  text = read_txt(data)
            else: continue
            for ch in chunk_text(text):
                if ch.strip():
                    all_chunks.append(ch)
                    metas.append({"source": fn})
        except Exception:
            continue
    dim = EMBEDDER.get_sentence_embedding_dimension()
    index = faiss.IndexFlatIP(dim)
    if all_chunks:
        embs = EMBEDDER.encode(all_chunks, convert_to_numpy=True, normalize_embeddings=True).astype("float32")
        index.add(embs)
    save_rag(index, all_chunks, metas)
    return index, all_chunks, metas

def ensure_rag():
    idx, ch, mt = load_rag()
    if idx is None: idx, ch, mt = rebuild_rag_from_docs()
    return idx, ch, mt

def rag_search(query: str, index, chunks, metas, top_k=5, min_score=0.2):
    if index is None or len(chunks) == 0: return []
    q = EMBEDDER.encode([query], convert_to_numpy=True, normalize_embeddings=True).astype("float32")
    D, I = index.search(q, top_k)
    hits = []
    for rank, (idx, score) in enumerate(zip(I[0], D[0]), start=1):
        if 0 <= idx < len(chunks) and float(score) >= min_score:
            hits.append({"rank": rank, "score": float(score), "text": chunks[idx], "meta": metas[idx]})
    return hits

# ---------- LLMs ----------
def ask_llm(prompt: str, system: str = None, model: str = DEFAULT_OLLAMA_MODEL, max_new_tokens=600) -> str:
    # Prefer Ollama for quality/latency (local/private)
    if OLLAMA_AVAILABLE:
        try:
            messages = []
            if system: messages.append({"role": "system", "content": system})
            messages.append({"role": "user", "content": prompt})
            out = ollama.chat(model=model, messages=messages)
            return out["message"]["content"]
        except Exception:
            pass
    # Fallback HF pipeline (smaller/less capable but works offline)
    if HF_AVAILABLE:
        text = (f"System: {system}\n" if system else "") + f"User: {prompt}\nAssistant:"
        out = hf_pipe(text, max_new_tokens=max_new_tokens, do_sample=False)[0]["generated_text"]
        return out.strip()
    return "No LLM backend available. Start Ollama or install transformers."

# ---------- Calculator ----------
def safe_eval(expr: str) -> str:
    allowed = {
        "abs": abs, "round": round, "min": min, "max": max, "pow": pow,
        "math": math, "sin": math.sin, "cos": math.cos, "tan": math.tan,
        "sqrt": math.sqrt, "log": math.log, "log10": math.log10, "pi": math.pi, "e": math.e
    }
    if not re.fullmatch(r"[0-9\.\s\+\-\*\/\%\(\)a-zA-Z_]+", expr):
        return "Error: invalid characters."
    try:
        return str(eval(expr, {"__builtins__": None}, allowed))
    except Exception as e:
        return f"Error: {e}"

def looks_like_math(s: str) -> bool:
    return bool(re.fullmatch(r"[0-9\.\s\+\-\*\/\%\(\)]+", s.strip()))

# ---------- Web search (summarize, no link-dump) ----------
def web_search_summarize(query: str, k=3) -> Tuple[str, List[Dict]]:
    results = []
    with DDGS() as ddgs:
        for r in ddgs.text(query, max_results=k):
            url = r.get("href") or r.get("url")
            title = r.get("title") or "(no title)"
            body = r.get("body") or ""
            if not url: continue
            snippet = body
            # Try to fetch clean text for better summaries
            try:
                resp = requests.get(url, timeout=10, headers={"User-Agent": "Mozilla/5.0"})
                doc = Document(resp.text)
                clean = BeautifulSoup(doc.summary(), "lxml").get_text()
                snippet = clean[:1500] if clean else snippet
            except Exception:
                pass
            results.append({"title": title, "url": url, "snippet": snippet})

    if not results:
        return "No reliable results found.", []

    combined = "\n\n".join([f"[W{i+1}] {r['title']}\n{r['snippet']}" for i, r in enumerate(results)])
    sys = ("You are a precise assistant. Use the evidence to answer naturally. "
           "Cite inline like [W1]/[W2] only when needed. Avoid listing links.")
    prompt = f"Query: {query}\n\nEvidence:\n{combined}\n\nWrite a concise, authoritative answer."
    answer = ask_llm(prompt, system=sys, max_new_tokens=600)
    return answer, results

# ---------- Routing & Answer composition ----------
ASSISTANT_SYSTEM = (
    "You are a helpful, professional assistant. "
    "First, answer from your own knowledge and reasoning. "
    "If document or web context is present, incorporate it and cite [D#]/[W#] only when you rely on it. "
    "If information is insufficient, say so briefly and suggest the next step."
)

def build_context_block(doc_hits: List[Dict], web_answer_text: str | None) -> str:
    parts = []
    for i, h in enumerate(doc_hits, start=1):
        src = h['meta'].get('source', '(doc)')
        parts.append(f"[D{i}] {src}\n{h['text']}")
    if web_answer_text:
        parts.append(web_answer_text)
    return "\n\n---\n\n".join(parts) if parts else "(no external context)"

def agent_answer(user_msg: str, use_docs=True, use_web=False, rag_k=4, web_k=3, model=DEFAULT_OLLAMA_MODEL):
    # 1) Calculator fast-path
    if looks_like_math(user_msg):
        return f"🧮 `{user_msg}` = **{safe_eval(user_msg)}**"

    # 2) Gather optional context
    doc_hits: List[Dict] = []
    web_text, web_sources = None, []

    if use_docs:
        index, chunks, metas = ensure_rag()
        if chunks:
            doc_hits = rag_search(user_msg, index, chunks, metas, top_k=rag_k, min_score=0.25)

    if use_web:
        web_text, web_sources = web_search_summarize(user_msg, k=web_k)

    # 3) Compose final prompt (reasoning-first, then blend context)
    context_block = build_context_block(doc_hits, web_text)
    prompt = (
        f"CONTEXT:\n{context_block}\n\n"
        f"USER QUESTION: {user_msg}\n\n"
        "INSTRUCTIONS:\n"
        "- Provide a direct, well-structured answer in natural language.\n"
        "- Use [D#]/[W#] citations only when you directly rely on those sources.\n"
        "- If no context is helpful, answer from general knowledge.\n"
        "- Keep it concise but complete."
    )
    answer = ask_llm(prompt, system=ASSISTANT_SYSTEM, model=model, max_new_tokens=700)

    # 4) Optional Sources block (for transparency)
    sources_lines = []
    for i, h in enumerate(doc_hits, start=1):
        sources_lines.append(f"- [D{i}] {h['meta'].get('source','(doc)')}")
    for i, r in enumerate(web_sources, start=1):
        sources_lines.append(f"- [W{i}] {r['title']} — {r['url']}")
    if sources_lines:
        answer = answer + "\n\n---\n**Sources**\n" + "\n".join(sources_lines)

    return answer

# ============================================================
# STREAMLIT (ChatGPT-like UI)
# ============================================================
def run_streamlit():
    import streamlit as st
    st.set_page_config(page_title=APP_TITLE, page_icon="🤖", layout="wide")
    st.title("🤖 Pro AI Agent")
    st.caption("Chat • Reasoning-first • RAG on your files • Web-aware • Calculator")

    with st.sidebar:
        st.subheader("Settings")
        st.write(f"Ollama: **{OLLAMA_AVAILABLE}** | HF fallback: **{HF_AVAILABLE}**")
        model = st.text_input("Ollama model", value=DEFAULT_OLLAMA_MODEL)
        use_docs = st.toggle("Use documents (RAG)", True)
        use_web = st.toggle("Use web", False)
        rag_k = st.slider("Doc passages", 1, 8, 4)
        web_k = st.slider("Web pages to summarize", 1, 5, 3)

        st.divider()
        st.subheader("Upload files (.txt/.pdf/.docx)")
        files = st.file_uploader("Add to knowledge base", type=["txt","pdf","docx"], accept_multiple_files=True)
        if files:
            for f in files:
                with open(os.path.join(DOCS_DIR, f.name), "wb") as out:
                    out.write(f.read())
            st.success("Saved. Click Rebuild index to include them.")
        if st.button("🔁 Rebuild index now"):
            with st.spinner("Indexing..."):
                rebuild_rag_from_docs()
            st.success("RAG index rebuilt.")

    if "history" not in st.session_state:
        st.session_state.history = []

    for m in st.session_state.history:
        with st.chat_message(m["role"]):
            st.markdown(m["content"])

    user_msg = st.chat_input("Ask anything…")
    if user_msg:
        st.session_state.history.append({"role":"user","content": user_msg})
        with st.chat_message("user"):
            st.markdown(user_msg)

        try:
            with st.spinner("Thinking…"):
                out = agent_answer(user_msg, use_docs=use_docs, use_web=use_web, rag_k=rag_k, web_k=web_k, model=model)
        except Exception:
            out = f"⚠️ Error:\n```\n{traceback.format_exc()}\n```"

        st.session_state.history.append({"role":"assistant","content": out})
        with st.chat_message("assistant"):
            st.markdown(out)

# ============================================================
# GRADIO (lightweight)
# ============================================================
def run_gradio():
    import gradio as gr

    def chat_fn(message, history, use_docs=True, use_web=False, rag_k=4, web_k=3, model=DEFAULT_OLLAMA_MODEL):
        try:
            reply = agent_answer(str(message), use_docs=bool(use_docs), use_web=bool(use_web),
                                 rag_k=int(rag_k), web_k=int(web_k), model=model)
        except Exception:
            reply = f"Error:\n{traceback.format_exc()}"
        history = (history or []) + [[message, reply]]
        return history, ""

    with gr.Blocks(title="Pro AI Agent") as demo:
        gr.Markdown("## 🤖 Pro AI Agent — Chat + RAG + Web")
        with gr.Row():
            use_docs = gr.Checkbox(True, label="Use documents (RAG)")
            use_web  = gr.Checkbox(False, label="Use web")
            rag_k    = gr.Slider(1, 8, value=4, step=1, label="Doc passages")
            web_k    = gr.Slider(1, 5, value=3, step=1, label="Web pages")
            model    = gr.Textbox(value=DEFAULT_OLLAMA_MODEL, label="Ollama model")

        chatbot = gr.Chatbot(height=520)
        msg = gr.Textbox(placeholder="Ask anything…")
        clear = gr.Button("Clear")

        msg.submit(chat_fn, [msg, chatbot, use_docs, use_web, rag_k, web_k, model], [chatbot, msg])
        clear.click(lambda: [], None, chatbot)

        gr.Markdown("> Upload files into the `docs/` folder and rebuild index in Streamlit if needed.")

    demo.launch()

# ============================================================
# MAIN
# ============================================================
if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1 and sys.argv[1].lower() == "gradio":
        run_gradio()
    else:
        run_streamlit()
