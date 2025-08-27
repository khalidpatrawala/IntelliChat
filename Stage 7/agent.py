import os, io, re, math, json, pickle, asyncio, traceback
from typing import List, Dict, Tuple, Optional

import numpy as np
import streamlit as st
import faiss

# --- Web search & fetch ---
import requests
from bs4 import BeautifulSoup
from readability import Document
try:
    from ddgs import DDGS          # preferred package
except Exception:
    from duckduckgo_search import DDGS  # fallback

import httpx  # for async fetches

# --- Files ---
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

# --- Embeddings (CPU only to prevent meta-tensor issue) ---
from sentence_transformers import SentenceTransformer

# --- LLM backends ---
DEFAULT_OLLAMA_MODEL = "llama3.1:8b"
OLLAMA_AVAILABLE = False
HF_AVAILABLE = False
try:
    import ollama
    OLLAMA_AVAILABLE = True
except Exception:
    OLLAMA_AVAILABLE = False

try:
    from transformers import pipeline
    HF_AVAILABLE = True
except Exception:
    HF_AVAILABLE = False

# ======================= CONFIG & PATHS =======================
APP_TITLE = "CognifyAI — Pro Chat Agent"
BASE_DIR = os.path.abspath(".")
DOCS_DIR = os.path.join(BASE_DIR, "docs")
os.makedirs(DOCS_DIR, exist_ok=True)

INDEX_FILE  = os.path.join(BASE_DIR, "faiss.index")
CHUNKS_FILE = os.path.join(BASE_DIR, "chunks.pkl")
METAS_FILE  = os.path.join(BASE_DIR, "metas.pkl")

EMBED_MODEL = "all-MiniLM-L6-v2"
EMBED_BATCH = 64             # keep small for low RAM
CHUNK_SIZE  = 1000           # chars
CHUNK_OVERL = 200            # chars
MAX_FILE_MB = 15             # hard limit per file
MAX_PDF_PAGES = 50           # avoid giant PDFs

# ======================= CACHED RESOURCES =====================
@st.cache_resource(show_spinner=False)
def get_embedder():
    # Force CPU to avoid meta-tensor crash on some Windows installs
    return SentenceTransformer(EMBED_MODEL, device="cpu")

@st.cache_resource(show_spinner=False)
def load_or_init_index() -> faiss.Index:
    dim = get_embedder().get_sentence_embedding_dimension()
    if os.path.exists(INDEX_FILE) and os.path.exists(CHUNKS_FILE) and os.path.exists(METAS_FILE):
        try:
            return faiss.read_index(INDEX_FILE)
        except Exception:
            pass
    # empty inner-product index
    return faiss.IndexFlatIP(dim)

@st.cache_resource(show_spinner=False)
def lazy_hf_pipeline():
    # Load only when user explicitly enables HF fallback
    from transformers import pipeline
    return pipeline("text2text-generation", model="google/flan-t5-large")

@st.cache_data(show_spinner=False)
def load_meta() -> Tuple[List[str], List[dict]]:
    if os.path.exists(CHUNKS_FILE) and os.path.exists(METAS_FILE):
        try:
            with open(CHUNKS_FILE, "rb") as f: chunks = pickle.load(f)
            with open(METAS_FILE, "rb") as f: metas  = pickle.load(f)
            return chunks, metas
        except Exception:
            pass
    return [], []

def persist_index(index: faiss.Index, chunks: List[str], metas: List[dict]):
    faiss.write_index(index, INDEX_FILE)
    with open(CHUNKS_FILE, "wb") as f: pickle.dump(chunks, f)
    with open(METAS_FILE, "wb") as f: pickle.dump(metas, f)
    # bust caches
    load_or_init_index.clear()
    load_meta.clear()

# ======================= UTILITIES ============================
def chunk_text(text: str, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERL) -> List[str]:
    text = text or ""
    chunks, n, start = [], len(text), 0
    while start < n:
        end = min(n, start + chunk_size)
        ck = text[start:end]
        if ck.strip():
            chunks.append(ck)
        if end == n: break
        start = max(0, end - overlap)
    return chunks

def read_txt(b: bytes) -> str:
    return b.decode("utf-8", errors="ignore")

def read_pdf(b: bytes) -> str:
    buf = io.BytesIO(b)
    reader = PdfReader(buf)
    texts = []
    for i, p in enumerate(reader.pages):
        if i >= MAX_PDF_PAGES:
            texts.append(f"\n[Truncated after {MAX_PDF_PAGES} pages to save memory]")
            break
        texts.append(p.extract_text() or "")
    return "\n".join(texts)

def read_docx(b: bytes) -> str:
    buf = io.BytesIO(b)
    doc = DocxDocument(buf)
    return "\n".join(p.text for p in doc.paragraphs)

def file_too_big(upload) -> bool:
    return (getattr(upload, "size", 0) or 0) > MAX_FILE_MB * 1024 * 1024

# ======================= RAG OPS ==============================
def rebuild_from_docs() -> Tuple[faiss.Index, List[str], List[dict]]:
    embedder = get_embedder()
    dim = embedder.get_sentence_embedding_dimension()
    index = faiss.IndexFlatIP(dim)
    chunks, metas = [], []

    # iterate docs folder, stream in batches
    for fn in os.listdir(DOCS_DIR):
        path = os.path.join(DOCS_DIR, fn)
        if not os.path.isfile(path): continue
        try:
            with open(path, "rb") as f:
                data = f.read()
            lo = fn.lower()
            if lo.endswith(".pdf"):
                text = read_pdf(data)
            elif lo.endswith(".docx"):
                text = read_docx(data)
            elif lo.endswith(".txt"):
                text = read_txt(data)
            else:
                continue

            new_chunks = chunk_text(text)
            if not new_chunks: continue

            # embed in small batches
            for i in range(0, len(new_chunks), EMBED_BATCH):
                batch = new_chunks[i:i+EMBED_BATCH]
                embs = embedder.encode(batch, convert_to_numpy=True, normalize_embeddings=True).astype("float32")
                index.add(embs)
                chunks.extend(batch)
                metas.extend([{"source": fn}] * len(batch))

        except Exception:
            continue

    persist_index(index, chunks, metas)
    return index, chunks, metas

def ensure_index_and_meta() -> Tuple[faiss.Index, List[str], List[dict]]:
    idx = load_or_init_index()
    chunks, metas = load_meta()
    return idx, chunks, metas

def rag_search(query: str, top_k=4, min_score=0.25) -> List[Dict]:
    index, chunks, metas = ensure_index_and_meta()
    if index is None or len(chunks) == 0:
        return []
    embedder = get_embedder()
    q = embedder.encode([query], convert_to_numpy=True, normalize_embeddings=True).astype("float32")
    D, I = index.search(q, top_k)
    hits = []
    for rank, (idx, score) in enumerate(zip(I[0], D[0]), start=1):
        if 0 <= idx < len(chunks) and float(score) >= min_score:
            hits.append({"rank": rank, "score": float(score), "text": chunks[idx], "meta": metas[idx]})
    return hits

# ======================= LLM WRAPPERS =========================
def stream_ollama(messages, model=DEFAULT_OLLAMA_MODEL):
    # Stream tokens (very memory-friendly)
    for part in ollama.chat(model=model, messages=messages, stream=True):
        yield part["message"]["content"]

def ask_llm(prompt: str, system: Optional[str], model: str, stream: bool):
    messages = []
    if system:
        messages.append({"role": "system", "content": system})
    messages.append({"role": "user", "content": prompt})

    if OLLAMA_AVAILABLE:
        if stream:
            return stream_ollama(messages, model=model)
        else:
            out = ollama.chat(model=model, messages=messages)
            return out["message"]["content"]

    if HF_AVAILABLE:
        # non-stream fallback; keep it optional to avoid OOM
        pipe = lazy_hf_pipeline()
        text = (f"System: {system}\n" if system else "") + f"User: {prompt}\nAssistant:"
        out = pipe(text, max_new_tokens=600, do_sample=False)[0]["generated_text"]
        return out

    return "No LLM backend available. Please start Ollama or enable HF fallback."

# ======================= WEB SEARCH (ASYNC) ===================
async def fetch_clean(session: httpx.AsyncClient, url: str) -> str:
    try:
        r = await session.get(url, timeout=10)
        r.raise_for_status()
        doc = Document(r.text)
        clean = BeautifulSoup(doc.summary(), "lxml").get_text()
        return clean[:1500] if clean else ""
    except Exception:
        return ""

async def gather_pages(urls: List[str], max_concurrency=5) -> List[str]:
    sem = asyncio.Semaphore(max_concurrency)
    texts = []

    async with httpx.AsyncClient(headers={"User-Agent": "Mozilla/5.0"}) as session:
        async def fetch_with_sem(u):
            async with sem:
                return await fetch_clean(session, u)
        tasks = [fetch_with_sem(u) for u in urls]
        texts = await asyncio.gather(*tasks, return_exceptions=False)

    return texts

def web_search_summarize(query: str, k=3, async_fetch=True) -> Tuple[str, List[Dict]]:
    # get top search results (sync)
    results = []
    with DDGS() as ddgs:
        for r in ddgs.text(query, max_results=k):
            url = r.get("href") or r.get("url")
            title = r.get("title") or "(no title)"
            body = r.get("body") or ""
            if url:
                results.append({"title": title, "url": url, "snippet": body})

    if not results:
        return "No reliable results found.", []

    # fetch pages (async or sync fallback)
    urls = [r["url"] for r in results]
    if async_fetch:
        try:
            texts = asyncio.run(gather_pages(urls))
        except RuntimeError:
            # if already in event loop, fallback to sync
            texts = []
            for u in urls:
                try:
                    resp = requests.get(u, timeout=10, headers={"User-Agent": "Mozilla/5.0"})
                    doc = Document(resp.text)
                    texts.append(BeautifulSoup(doc.summary(), "lxml").get_text()[:1500])
                except Exception:
                    texts.append(results[len(texts)]["snippet"])
    else:
        texts = []
        for u in urls:
            try:
                resp = requests.get(u, timeout=10, headers={"User-Agent": "Mozilla/5.0"})
                doc = Document(resp.text)
                texts.append(BeautifulSoup(doc.summary(), "lxml").get_text()[:1500])
            except Exception:
                texts.append(results[len(texts)]["snippet"])

    # replace snippets with fetched clean text when available
    for i, t in enumerate(texts):
        if t: results[i]["snippet"] = t

    # summarize via LLM (concise; avoid link-dumps)
    evidence = "\n\n".join([f"[W{i+1}] {r['title']}\n{r['snippet']}" for i, r in enumerate(results)])
    sys = ("You are a precise assistant. Use the evidence to answer naturally. "
           "Cite inline like [W1]/[W2] only when you rely on a claim from that source. Keep it concise.")
    prompt = f"Query: {query}\n\nEvidence:\n{evidence}\n\nWrite a direct, authoritative answer."
    answer = ask_llm(prompt, system=sys, model=st.session_state.get("model", DEFAULT_OLLAMA_MODEL), stream=False)
    return answer if isinstance(answer, str) else "".join(answer), results

# ======================= ROUTER ===============================
ASSISTANT_SYSTEM = (
    "You are a helpful, professional assistant. "
    "Answer from reasoning first. If document or web context is present, use it and cite [D#]/[W#] only when relied upon. "
    "Be concise but complete. If uncertain, say what else is needed."
)

def looks_like_math(s: str) -> bool:
    return bool(re.fullmatch(r"[0-9\.\s\+\-\*\/\%\(\)]+", s.strip()))

def safe_eval(expr: str) -> str:
    allowed = {
        "abs": abs, "round": round, "min": min, "max": max, "pow": pow,
        "math": math, "sin": math.sin, "cos": math.cos, "tan": math.tan,
        "sqrt": math.sqrt, "log": math.log, "log10": math.log10, "pi": math.pi, "e": math.e
    }
    if not re.fullmatch(r"[0-9\.\s\+\-\*\/\%\(\)a-zA-Z_]+", expr):
        return "Error: invalid characters."
    try:
        return str(eval(expr, {"__builtins__": None}, allowed))
    except Exception as e:
        return f"Error: {e}"

def compose_prompt(user_msg: str, doc_hits: List[Dict], web_text: Optional[str]) -> str:
    parts = []
    for i, h in enumerate(doc_hits, start=1):
        src = h['meta'].get('source', '(doc)')
        parts.append(f"[D{i}] {src}\n{h['text']}")
    if web_text:
        parts.append(web_text)
    context = "\n\n---\n\n".join(parts) if parts else "(no external context)"
    return (
        f"CONTEXT:\n{context}\n\n"
        f"USER: {user_msg}\n\n"
        "INSTRUCTIONS:\n"
        "- Provide a direct, well-structured answer.\n"
        "- Only cite [D#]/[W#] when directly relying on that source.\n"
        "- Keep it concise but complete."
    )

def agent_answer(user_msg: str, use_docs=True, use_web=False, rag_k=4, web_k=3, stream=True):
    # Fast-path calculator to avoid LLM call
    if looks_like_math(user_msg):
        return f"🧮 `{user_msg}` = **{safe_eval(user_msg)}**", None

    doc_hits = rag_search(user_msg, top_k=rag_k) if use_docs else []
    web_text, web_sources = (None, [])
    if use_web:
        web_text, web_sources = web_search_summarize(user_msg, k=web_k, async_fetch=True)

    prompt = compose_prompt(user_msg, doc_hits, web_text)
    model = st.session_state.get("model", DEFAULT_OLLAMA_MODEL)

    gen = ask_llm(prompt, system=ASSISTANT_SYSTEM, model=model, stream=stream)

    # Build optional sources block
    sources_lines = []
    for i, h in enumerate(doc_hits, start=1):
        sources_lines.append(f"- [D{i}] {h['meta'].get('source','(doc)')}")
    for i, r in enumerate(web_sources, start=1):
        sources_lines.append(f"- [W{i}] {r['title']} — {r['url']}")

    return gen, ("\n\n---\n**Sources**\n" + "\n".join(sources_lines) if sources_lines else None)

# ======================= STREAMLIT UI =========================
st.set_page_config(page_title=APP_TITLE, page_icon="🤖", layout="wide")
st.title("🤖 IntelliChat")
st.caption("Chat • Reasoning-first • RAG on your files • Web-aware • Low-memory streaming")

with st.sidebar:
    st.subheader("Settings")
    st.write(f"Ollama: **{OLLAMA_AVAILABLE}**  |  HF fallback: **{HF_AVAILABLE}**")
    st.session_state.model = st.text_input("Ollama model", value=DEFAULT_OLLAMA_MODEL)

    use_docs = st.toggle("Use documents (RAG)", True)
    use_web  = st.toggle("Use web search", False)
    rag_k    = st.slider("Doc passages", 1, 8, 4)
    web_k    = st.slider("Web pages", 1, 5, 3)

    st.divider()
    st.subheader("Add documents (.txt/.pdf/.docx)")
    uploads = st.file_uploader("Files (each ≤ {} MB)".format(MAX_FILE_MB), accept_multiple_files=True, type=["txt","pdf","docx"])
    if uploads:
        saved_any = False
        for up in uploads:
            if file_too_big(up):
                st.warning(f"Skipped `{up.name}` (>{MAX_FILE_MB} MB).")
                continue
            path = os.path.join(DOCS_DIR, up.name)
            with open(path, "wb") as f:
                f.write(up.read())
            saved_any = True
        if saved_any:
            st.success("Saved. Click **Rebuild index** below.")
    if st.button("🔁 Rebuild index"):
        with st.spinner("Indexing (batched, memory-safe)..."):
            rebuild_from_docs()
        st.success("RAG index rebuilt.")

# Initialize chat memory
if "chat" not in st.session_state:
    st.session_state.chat = []

# Render history
for msg in st.session_state.chat:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])

# Input
if user_msg := st.chat_input("Ask anything…"):
    st.session_state.chat.append({"role":"user","content": user_msg})
    with st.chat_message("user"):
        st.markdown(user_msg)

    try:
        gen, sources_block = agent_answer(user_msg, use_docs=use_docs, use_web=use_web, rag_k=rag_k, web_k=web_k, stream=True)

        with st.chat_message("assistant"):
            if isinstance(gen, str):
                st.markdown(gen)
                full = gen
            else:
                # Stream tokens from Ollama
                chunks = []
                def streamer():
                    for t in gen:
                        chunks.append(t)
                        yield t
                full = st.write_stream(streamer)

            if sources_block:
                st.markdown(sources_block)

        # persist full assistant message
        st.session_state.chat.append({"role":"assistant","content": full + (sources_block or "")})

    except Exception:
        err = f"⚠️ Error:\n```\n{traceback.format_exc()}\n```"
        with st.chat_message("assistant"):
            st.markdown(err)
        st.session_state.chat.append({"role":"assistant","content": err})
