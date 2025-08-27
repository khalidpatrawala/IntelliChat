import os
import re
import io
import json
import time
import math
import ollama
import requests
import numpy as np
from bs4 import BeautifulSoup
from readability import Document
from ddgs import DDGS
from rich.console import Console
from rich.markdown import Markdown
import docx
import PyPDF2
import faiss
from sentence_transformers import SentenceTransformer
from contextlib import redirect_stdout

# ======================
# Config and constants
# ======================
console = Console()
HISTORY_FILE = "history.json"
DOCS_FOLDER = "docs"
NOTES_FOLDER = "notes"
REMINDERS_FILE = "reminders.json"
EMBEDDINGS_FILE = "faiss_index.index"
DOC_TEXTS_FILE = "doc_texts.json"
MODEL_NAME = "llama3.1:8b"   # change to 'mistral:7b' if needed

os.makedirs(DOCS_FOLDER, exist_ok=True)
os.makedirs(NOTES_FOLDER, exist_ok=True)

# ======================
# Memory (Stage 2)
# ======================
def load_history():
    if os.path.exists(HISTORY_FILE):
        with open(HISTORY_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return []

def save_history(history):
    with open(HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump(history, f, indent=2)

# ======================
# Web Search (Stage 3)
# ======================
def web_search(query, max_results=3):
    """DuckDuckGo search + readable summaries"""
    results = []
    with DDGS() as ddgs:
        for r in ddgs.text(query, max_results=max_results):
            url = r.get("href")
            if not url:
                continue
            try:
                resp = requests.get(url, timeout=12, headers={"User-Agent":"Mozilla/5.0"})
                doc = Document(resp.text)
                title = doc.title() or "(no title)"
                text = BeautifulSoup(doc.summary(), "lxml").get_text()
                results.append(f"{title}\n{text[:600]}... [source: {url}]")
            except Exception as e:
                results.append(f"Could not fetch {url} ({e})")
    return results

# ======================
# Document reading (Stage 4)
# ======================
def read_text_file(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def read_docx_file(path):
    d = docx.Document(path)
    return "\n".join(p.text for p in d.paragraphs)

def read_pdf_file(path):
    text = ""
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() or ""
    return text

def load_documents(folder=DOCS_FOLDER):
    texts = []
    for fn in os.listdir(folder):
        path = os.path.join(folder, fn)
        if not os.path.isfile(path):
            continue
        if fn.lower().endswith(".txt"):
            texts.append(read_text_file(path))
        elif fn.lower().endswith(".docx"):
            texts.append(read_docx_file(path))
        elif fn.lower().endswith(".pdf"):
            texts.append(read_pdf_file(path))
    return texts

# ======================
# Persistent RAG (Stage 5)
# ======================
embedder = SentenceTransformer("all-MiniLM-L6-v2")

def build_or_load_faiss(docs_folder=DOCS_FOLDER):
    # Load existing
    if os.path.exists(EMBEDDINGS_FILE) and os.path.exists(DOC_TEXTS_FILE):
        try:
            index = faiss.read_index(EMBEDDINGS_FILE)
            with open(DOC_TEXTS_FILE, "r", encoding="utf-8") as f:
                texts = json.load(f)
            return index, texts
        except Exception:
            pass  # fall through to rebuild

    # Build fresh
    texts = load_documents(docs_folder)
    if not texts:
        return None, []

    embeddings = embedder.encode(texts, convert_to_numpy=True, normalize_embeddings=True)
    dim = embeddings.shape[1]
    index = faiss.IndexFlatIP(dim)
    index.add(embeddings.astype("float32"))

    faiss.write_index(index, EMBEDDINGS_FILE)
    with open(DOC_TEXTS_FILE, "w", encoding="utf-8") as f:
        json.dump(texts, f, indent=2)
    return index, texts

def search_faiss(query, index, texts, top_k=3):
    if index is None or not texts:
        return []
    q = embedder.encode([query], convert_to_numpy=True, normalize_embeddings=True).astype("float32")
    D, I = index.search(q, top_k)
    hits = []
    for idx, score in zip(I[0], D[0]):
        if 0 <= idx < len(texts):
            hits.append((float(score), texts[idx][:1500]))
    return hits

# ======================
# Tools (Stage 6)
# ======================
# 1) Calculator (safe-ish eval limited to math)
def safe_eval(expr):
    allowed_names = {
        "abs": abs, "round": round, "min": min, "max": max, "pow": pow,
        "math": math,
        # common math symbols
        "sin": math.sin, "cos": math.cos, "tan": math.tan,
        "sqrt": math.sqrt, "log": math.log, "log10": math.log10, "pi": math.pi, "e": math.e
    }
    # allow digits, operators, letters, parentheses, dot, comma, spaces
    if not re.fullmatch(r"[0-9\.\,\s\+\-\*\/\%\(\)a-zA-Z_]+", expr):
        return "Error: invalid characters in expression."
    try:
        return str(eval(expr, {"__builtins__": {}}, allowed_names))
    except Exception as e:
        return f"Error: {e}"

# 2) Python code execution (very constrained)
def run_python(code):
    """
    Executes tiny Python snippets with no builtins; captures print output.
    WARNING: still run only code you trust. This is a minimal sandbox.
    """
    buf = io.StringIO()
    # Very limited environment
    env = {"__builtins__": {}}
    # Provide a safe 'print' that writes to buffer
    def _print(*args, **kwargs):
        print(*args, **kwargs, file=buf)
    env["print"] = _print
    # Optionally expose math
    env["math"] = math
    locals_dict = {}
    try:
        with redirect_stdout(buf):
            exec(code, env, locals_dict)
        out = buf.getvalue().strip()
        if out:
            return out
        # If no prints, return local variables for visibility
        if locals_dict:
            try:
                return json.dumps({k: v for k, v in locals_dict.items() if not k.startswith("_")}, default=str, indent=2)
            except Exception:
                return str(locals_dict)
        return "(no output)"
    except Exception as e:
        return f"Error: {e}"

# 3) Notes writer
def sanitize_filename(name):
    name = name.strip().replace(" ", "_")
    return re.sub(r"[^a-zA-Z0-9_\-\.]", "", name) or f"note_{int(time.time())}.txt"

def save_note(filename, content):
    fname = sanitize_filename(filename)
    path = os.path.join(NOTES_FOLDER, fname)
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    return f"Saved note to {path}"

# 4) Reminders / To-Dos (simple list)
def load_reminders():
    if os.path.exists(REMINDERS_FILE):
        with open(REMINDERS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return []

def save_reminders(items):
    with open(REMINDERS_FILE, "w", encoding="utf-8") as f:
        json.dump(items, f, indent=2)

def add_reminder(text):
    items = load_reminders()
    items.append({"text": text, "ts": int(time.time())})
    save_reminders(items)
    return "Reminder saved."

def list_reminders():
    items = load_reminders()
    if not items:
        return "No reminders yet."
    lines = []
    for i, r in enumerate(items, 1):
        t = time.strftime("%Y-%m-%d %H:%M", time.localtime(r["ts"]))
        lines.append(f"{i}. {r['text']}  —  added {t}")
    return "\n".join(lines)

# ======================
# Agent core
# ======================
def agent_reply(history, user_msg):
    msg = user_msg.strip()
    lo = msg.lower()

    # --- Tools / commands ---
    if lo.startswith("help"):
        return (
            "**Commands**\n"
            "- `search: <query>` → web search\n"
            "- `rag: <query>` → answer from your docs (persistent RAG)\n"
            "- `docs` → quick dump/preview of docs\n"
            "- `calc: <expression>` → calculator (e.g., calc: 2*(5+3))\n"
            "- `run: <python>` → run tiny Python code (print, math)\n"
            "- `note:<filename>:<content>` → save a note in notes/\n"
            "- `reminder:<text>` → add a reminder\n"
            "- `list reminders` → list all reminders\n"
            "- `exit` → quit"
        )

    # Web search
    if lo.startswith("search:"):
        query = msg.split("search:", 1)[1].strip()
        results = web_search(query)
        return "\n\n".join(results) if results else "No results found."

    # RAG: question over docs with persistent FAISS
    if lo.startswith("rag:"):
        query = msg.split("rag:", 1)[1].strip()
        index, texts = build_or_load_faiss()
        hits = search_faiss(query, index, texts, top_k=3)
        if not hits:
            return "No documents found for RAG. Put .txt/.pdf/.docx files in docs/."
        context = "\n\n---\n\n".join(f"[score {round(s,3)}]\n{snip}" for s, snip in hits)
        prompt = (
            "Use ONLY the following document excerpts to answer.\n"
            "Cite evidence briefly.\n\n"
            f"{context}\n\nQuestion: {query}"
        )
        response = ollama.chat(model=MODEL_NAME, messages=history + [{"role":"user","content":prompt}])
        reply = response["message"]["content"]
        history.append({"role":"assistant","content":reply})
        return reply

    # Docs quick preview
    if lo == "docs":
        texts = load_documents()
        if not texts:
            return "No documents found. Add .txt/.pdf/.docx files into the docs/ folder."
        joined = "\n\n---\n\n".join(t[:800] for t in texts)
        return f"Preview of your documents (truncated):\n\n{joined}"

    # Calculator
    if lo.startswith("calc:"):
        expr = msg.split("calc:", 1)[1].strip()
        return safe_eval(expr)

    # Python runner
    if lo.startswith("run:"):
        code = msg.split("run:", 1)[1].strip()
        return run_python(code)

    # Notes
    if lo.startswith("note:"):
        try:
            _, filename, content = msg.split(":", 2)
            return save_note(filename, content)
        except ValueError:
            return "Usage: note:<filename>:<content>"

    # Reminders
    if lo.startswith("reminder:"):
        text = msg.split("reminder:", 1)[1].strip()
        if not text:
            return "Usage: reminder:<text>"
        return add_reminder(text)

    if lo == "list reminders":
        return list_reminders()

    # --- Normal chat ---
    history.append({"role": "user", "content": msg})
    response = ollama.chat(model=MODEL_NAME, messages=history)
    reply = response["message"]["content"]
    history.append({"role": "assistant", "content": reply})
    return reply

# ======================
# Main loop
# ======================
def main():
    console.print("[bold green]AI Agent — Stage 6 (Tools + Persistent RAG + Web + Memory)[/bold green]")
    console.print("Type 'help' for commands. Type 'exit' to quit.\n")

    history = load_history()

    while True:
        try:
            user_msg = input("\nYou: ").strip()
        except (EOFError, KeyboardInterrupt):
            console.print("\n[red]Goodbye![/red]")
            save_history(history)
            break

        if not user_msg:
            continue
        if user_msg.lower() in {"exit", "quit", "bye"}:
            console.print("[red]Goodbye![/red]")
            save_history(history)
            break

        with console.status("Thinking..."):
            out = agent_reply(history, user_msg)

        console.print(Markdown(out))
        save_history(history)

if __name__ == "__main__":
    main()
