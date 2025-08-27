import os
import json
import ollama
from rich.console import Console
from rich.markdown import Markdown
from duckduckgo_search import DDGS
import requests
from bs4 import BeautifulSoup
from readability import Document
import docx
import PyPDF2
from sentence_transformers import SentenceTransformer
import faiss
import numpy as np

# ------------------------
# Config
# ------------------------
console = Console()
HISTORY_FILE = "history.json"
DOCS_FOLDER = "docs"
EMBEDDINGS_FILE = "faiss_index.index"
DOC_TEXTS_FILE = "doc_texts.json"

# ------------------------
# Memory functions
# ------------------------
def load_history():
    if os.path.exists(HISTORY_FILE):
        with open(HISTORY_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return []

def save_history(history):
    with open(HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump(history, f, indent=2)

# ------------------------
# Web search function
# ------------------------
def web_search(query, max_results=3):
    results = []
    with DDGS() as ddgs:
        for r in ddgs.text(query, max_results=max_results):
            url = r.get("href")
            if not url:
                continue
            try:
                resp = requests.get(url, timeout=10)
                doc = Document(resp.text)
                title = doc.title()
                summary = doc.summary()
                text = BeautifulSoup(summary, "lxml").get_text()
                results.append(f"{title}\n{text[:500]}... [source: {url}]")
            except Exception as e:
                results.append(f"Could not fetch {url} ({e})")
    return results

# ------------------------
# Document reading functions
# ------------------------
def read_text_file(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def read_docx_file(path):
    doc = docx.Document(path)
    return "\n".join([para.text for para in doc.paragraphs])

def read_pdf_file(path):
    text = ""
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() or ""
    return text

def load_documents(folder=DOCS_FOLDER):
    if not os.path.exists(folder):
        os.makedirs(folder)
        return []

    texts = []
    for fn in os.listdir(folder):
        path = os.path.join(folder, fn)
        if fn.endswith(".txt"):
            texts.append(read_text_file(path))
        elif fn.endswith(".docx"):
            texts.append(read_docx_file(path))
        elif fn.endswith(".pdf"):
            texts.append(read_pdf_file(path))
    return texts

# ------------------------
# RAG / FAISS functions
# ------------------------
embedder = SentenceTransformer("all-MiniLM-L6-v2")

def build_or_load_faiss(docs_folder=DOCS_FOLDER):
    # Load existing index
    if os.path.exists(EMBEDDINGS_FILE) and os.path.exists(DOC_TEXTS_FILE):
        index = faiss.read_index(EMBEDDINGS_FILE)
        with open(DOC_TEXTS_FILE, "r", encoding="utf-8") as f:
            texts = json.load(f)
        return index, texts

    # Build new index
    texts = load_documents(docs_folder)
    if not texts:
        return None, []

    embeddings = embedder.encode(texts, convert_to_numpy=True)
    index = faiss.IndexFlatL2(embeddings.shape[1])
    index.add(embeddings)

    # Save index & texts
    faiss.write_index(index, EMBEDDINGS_FILE)
    with open(DOC_TEXTS_FILE, "w", encoding="utf-8") as f:
        json.dump(texts, f, indent=2)

    return index, texts

def search_faiss(query, index, texts, top_k=1):
    if index is None or not texts:
        return ""
    q_emb = embedder.encode([query], convert_to_numpy=True)
    D, I = index.search(q_emb, top_k)
    return "\n\n".join([texts[i] for i in I[0]])

# ------------------------
# Agent reply function
# ------------------------
def agent_reply(history, user_msg):
    user_msg_lower = user_msg.lower().strip()

    # Web search
    if user_msg_lower.startswith("search:"):
        query = user_msg.split("search:",1)[1].strip()
        results = web_search(query)
        return "\n\n".join(results) if results else "No results found."

    # RAG / document question
    if user_msg_lower.startswith("rag:"):
        query = user_msg.split("rag:",1)[1].strip()
        index, texts = build_or_load_faiss()
        context = search_faiss(query, index, texts)
        if context:
            prompt = f"Answer the question using the following documents:\n\n{context}\n\nQuestion: {query}"
            response = ollama.chat(
                model="llama3.1:8b",
                messages=history + [{"role":"user","content":prompt}]
            )
            reply = response["message"]["content"]
            history.append({"role":"assistant","content":reply})
            return reply
        else:
            return "No documents found in RAG."

    # Normal chat
    history.append({"role": "user", "content": user_msg})
    response = ollama.chat(
        model="llama3.1:8b",
        messages=history
    )
    reply = response["message"]["content"]
    history.append({"role": "assistant", "content": reply})
    return reply

# ------------------------
# Main loop
# ------------------------
def main():
    console.print("[bold green]AI Agent: Stage 5 (Persistent RAG) Ready![/bold green]")
    console.print("Type 'exit' to quit, 'search: <query>' for web, 'rag: <query>' for document-based answers\n")

    history = load_history()

    while True:
        user_msg = input("\nYou: ").strip()
        if not user_msg:
            continue

        if user_msg.lower() in ["exit", "quit", "bye"]:
            console.print("[red]Goodbye![/red]")
            save_history(history)
            break

        with console.status("Thinking..."):
            out = agent_reply(history, user_msg)

        console.print(Markdown(out))
        save_history(history)

if __name__ == "__main__":
    main()
